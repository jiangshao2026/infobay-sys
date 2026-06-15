# -*- coding: utf-8 -*-
"""
将 Markdown 文件转换为 Word 文档 (.docx)
"""

import sys
import os

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.enum.text import WD_BREAK
except ImportError:
    print("正在安装 python-docx...")
    os.system('pip install python-docx')
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.enum.text import WD_BREAK


def set_chinese_font(run, size=11, bold=False):
    """设置中文字体"""
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size)
    if bold:
        run.bold = True


def add_title(doc, text, level=1):
    """添加标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    
    if level == 1:
        run.font.size = Pt(18)
        run.bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        run.font.color.rgb = RGBColor(0x1f, 0x38, 0x60)
    elif level == 2:
        run.font.size = Pt(15)
        run.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run.font.color.rgb = RGBColor(0x2c, 0x52, 0x80)
    elif level == 3:
        run.font.size = Pt(13)
        run.bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run.font.color.rgb = RGBColor(0x46, 0x82, 0xb4)
    elif level == 4:
        run.font.size = Pt(12)
        run.bold = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
    
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_paragraph(doc, text, size=11):
    """添加正文段落"""
    # 处理粗体文本 **text**
    parts = []
    i = 0
    while i < len(text):
        if text[i:i+2] == '**':
            # 查找结束标记
            end = text.find('**', i + 2)
            if end != -1:
                parts.append((text[i+2:end], True))
                i = end + 2
            else:
                parts.append((text[i:], False))
                break
        else:
            # 查找下一个 **
            next_bold = text.find('**', i)
            if next_bold != -1:
                parts.append((text[i:next_bold], False))
                i = next_bold
            else:
                parts.append((text[i:], False))
                break
    
    p = doc.add_paragraph()
    for part_text, is_bold in parts:
        run = p.add_run(part_text)
        set_chinese_font(run, size=size, bold=is_bold)
    p.paragraph_format.space_after = Pt(2)


def add_code_block(doc, text, size=9):
    """添加代码块"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # 设置背景色（浅色灰色背景）
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F5F5F5'
    })
    pPr.append(shd)
    
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)


def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # 填充表头
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        set_chinese_font(run, size=10, bold=True)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        # 设置表头背景色
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): '1F3860'
        })
        tcPr.append(shd)
        
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 填充数据行
    for idx, row in enumerate(rows):
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            # 处理单元格中的粗体
            cell_text_clean = cell_text.replace('**', '')
            row_cells[i].text = ''
            p = row_cells[i].paragraphs[0]
            run = p.add_run(cell_text_clean)
            set_chinese_font(run, size=10)
            
            # 偶数行添加背景色
            if idx % 2 == 1:
                tc = row_cells[i]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = tcPr.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'F0F4F8'
                })
                tcPr.append(shd)
            
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # 设置表格宽度
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(16.0 / len(headers))
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


def add_list_item(doc, text, level=0, size=11):
    """添加列表项"""
    # 处理粗体
    text_clean = text.replace('**', '')
    
    # 判断有序还是无序列表
    if text.strip().startswith(('- ', '* ', '+ ')) or (len(text.strip()) > 0 and text.strip()[0].isdigit() and text.strip()[1:3] in ['. ', ') ']):
        # 有序列表（简化处理：数字开头的视为有序）
        text_clean = text_clean.lstrip('-*+ ')
        if len(text_clean) > 0 and text_clean[0].isdigit():
            # 去除数字开头
            parts = text_clean.split(' ', 1)
            if len(parts) > 1:
                text_clean = parts[1]
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(level * 0.8)
    
    # 添加项目符号
    run_bullet = p.add_run('• ')
    set_chinese_font(run_bullet, size=size, bold=True)
    run_bullet.font.color.rgb = RGBColor(0x1f, 0x38, 0x60)
    
    run = p.add_run(text_clean.strip())
    set_chinese_font(run, size=size)
    p.paragraph_format.space_after = Pt(1)


def add_horizontal_rule(doc):
    """添加水平线"""
    p = doc.add_paragraph()
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_markdown(md_content):
    """解析 Markdown 内容为结构化数据"""
    lines = md_content.split('\n')
    elements = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 空行
        if not stripped:
            i += 1
            continue
        
        # 标题
        if stripped.startswith('#'):
            level = 0
            while level < len(stripped) and stripped[level] == '#':
                level += 1
            title_text = stripped[level:].strip()
            elements.append(('title', level, title_text))
            i += 1
            continue
        
        # 水平线
        if stripped == '---' or stripped == '***':
            elements.append(('hr',))
            i += 1
            continue
        
        # 代码块
        if stripped.startswith('```') or stripped.startswith('~~~'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith(('```', '~~~')):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束标记
            elements.append(('code', '\n'.join(code_lines)))
            continue
        
        # 表格检测
        if '|' in stripped and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if '|' in next_line and '---' in next_line:
                # 解析表格
                table_lines = [stripped]
                i += 1
                while i < len(lines) and '|' in lines[i].strip():
                    table_lines.append(lines[i].strip())
                    i += 1
                
                # 解析表格行
                rows = []
                for t_line in table_lines:
                    cells = [c.strip() for c in t_line.split('|')]
                    # 移除首尾空单元格（如果 | 在开头/结尾）
                    if cells and cells[0] == '':
                        cells = cells[1:]
                    if cells and cells[-1] == '':
                        cells = cells[:-1]
                    rows.append(cells)
                
                # 过滤掉分隔行
                valid_rows = []
                for row in rows:
                    if row and all(cell.strip() == '' or '---' in cell for cell in row):
                        continue
                    valid_rows.append(row)
                
                if valid_rows and len(valid_rows) > 0:
                    headers = valid_rows[0]
                    data_rows = valid_rows[1:]
                    elements.append(('table', headers, data_rows))
                continue
        
        # 列表项
        if stripped.startswith(('- ', '* ', '+ ')):
            elements.append(('list', stripped, 0))
            i += 1
            continue
        
        # 有序列表项（简化）
        if len(stripped) > 1 and stripped[0].isdigit():
            elements.append(('list', stripped, 0))
            i += 1
            continue
        
        # 普通段落
        elements.append(('paragraph', line))
        i += 1
    
    return elements


def convert_md_to_docx(md_file_path, docx_file_path):
    """将 Markdown 文件转换为 Word 文档"""
    
    # 读取 Markdown 文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 解析 Markdown 内容
    elements = parse_markdown(md_content)
    
    # 生成 Word 文档
    for element in elements:
        if element[0] == 'title':
            level, text = element[1], element[2]
            add_title(doc, text, level)
        
        elif element[0] == 'paragraph':
            text = element[1]
            add_paragraph(doc, text)
        
        elif element[0] == 'code':
            code_text = element[1]
            add_code_block(doc, code_text)
        
        elif element[0] == 'table':
            headers, rows = element[1], element[2]
            add_table(doc, headers, rows)
        
        elif element[0] == 'list':
            text = element[1]
            add_list_item(doc, text)
        
        elif element[0] == 'hr':
            add_horizontal_rule(doc)
    
    # 保存 Word 文档
    doc.save(docx_file_path)
    print(f"✅ Word 文档已生成: {docx_file_path}")


if __name__ == '__main__':
    # 获取脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    md_file = os.path.join(script_dir, '系统开发小结.md')
    docx_file = os.path.join(script_dir, '信佰监理服务管理系统-开发小结.docx')
    
    # 检查 Markdown 文件是否存在
    if not os.path.exists(md_file):
        print(f"❌ 找不到文件: {md_file}")
        sys.exit(1)
    
    print("正在读取: {}".format(md_file))
    print("正在生成 Word 文档...")
    
    convert_md_to_docx(md_file, docx_file)
